"""docx::generate_report — 生成 docx + CSV 审核报告。"""
from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .models import AuditIssue, dicts_to_issues

_SEV_COLOR = {
    "ERROR": RGBColor(0xCC, 0x00, 0x00),
    "WARNING": RGBColor(0xCC, 0x88, 0x00),
    "INFO": RGBColor(0x33, 0x66, 0x99),
}

_RULE_NAMES = {
    "AI_TRACE": "AI 生成痕迹",
    "HEADING_WITH_COMMENT": "标题含批注",
    "PARAGRAPH_WITHOUT_COMMENT": "段落缺批注",
    "TABLE_NAME_WITHOUT_TABLE": "有表名无表格",
    "PARAGRAPH_QUALITY": "段落语言质量",
}


def generate_csv_report(
    issues: list[AuditIssue], output_path: str, project: str = "DOC"
) -> str:
    csv_path = output_path.replace(".docx", ".csv")
    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["Key", "P", "类型", "问题", "修改建议", "位置"])
        for idx, issue in enumerate(issues, 1):
            key = f"{project}-{idx:03d}"
            location = (
                issue.context[:30].replace("\n", " ") if issue.context else ""
            )
            writer.writerow(
                [
                    key,
                    issue.priority,
                    issue.issue_type,
                    issue.message[:80],
                    (issue.suggestion or "")[:80],
                    location,
                ]
            )
    return csv_path


def generate_report(
    elements: list[dict],
    issues: list[AuditIssue],
    output_path: str,
    source_name: str = "",
) -> dict:
    doc = Document()
    title = doc.add_heading("文生文文档审核报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_n = sum(1 for e in elements if e["kind"] == "paragraph")
    h_n = sum(1 for e in elements if e["kind"] == "heading")
    t_n = sum(1 for e in elements if e["kind"] == "table")

    name = source_name or os.path.basename(output_path).replace(
        "_audit_report.docx", ""
    )
    doc.add_paragraph(f"源文件: {name}")
    doc.add_paragraph(f"段落: {p_n}  标题: {h_n}  表格: {t_n}  问题: {len(issues)}")
    doc.add_paragraph("")

    doc.add_heading("诊断明细", level=1)
    if issues:
        for issue in issues:
            p = doc.add_paragraph()
            r_sev = p.add_run(f"[{issue.severity}] ")
            r_sev.bold = True
            r_sev.font.color.rgb = _SEV_COLOR.get(
                issue.severity, RGBColor(0, 0, 0)
            )
            rule_name = _RULE_NAMES.get(issue.rule_id, issue.rule_id)
            p.add_run(f"({rule_name}) {issue.message}")

            if issue.context:
                cp = doc.add_paragraph()
                cr = cp.add_run(f"  上下文: {issue.context[:200]}")
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                cr.italic = True

            if issue.suggestion:
                sp = doc.add_paragraph()
                sr = sp.add_run(f"  建议: {issue.suggestion}")
                sr.font.size = Pt(10)
                sr.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    else:
        doc.add_paragraph("未发现问题。")

    doc.add_paragraph("")
    doc.add_heading("总结", level=1)
    errs = sum(1 for i in issues if i.severity == "ERROR")
    warns = sum(1 for i in issues if i.severity == "WARNING")

    if errs:
        doc.add_paragraph(
            f"发现 {errs} 个错误、{warns} 个警告。请优先处理 ERROR 级别条目。"
        )
    elif warns:
        doc.add_paragraph(f"无严重错误，有 {warns} 个警告建议改进。")
    else:
        doc.add_paragraph("未发现明显问题。")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    project_name = Path(output_path).parent.name or "DOC"
    csv_path = generate_csv_report(issues, output_path, project_name)

    return {
        "report_path": output_path,
        "csv_path": csv_path,
        "errors": errs,
        "warnings": warns,
        "total": len(issues),
    }


async def fn_generate_report(payload: dict) -> dict:
    """
    Function: docx::generate_report

    Input:
      {
        "elements": [...],
        "issues": [...],          # AuditIssue dicts
        "output_path": "...",
        "source_name": "..."      # optional
      }
    """
    elements = payload.get("elements") or []
    issues = dicts_to_issues(payload.get("issues") or [])
    output_path = payload.get("output_path")
    if not output_path:
        return {"error": "missing output_path"}
    source_name = payload.get("source_name", "")
    return generate_report(elements, issues, output_path, source_name)
